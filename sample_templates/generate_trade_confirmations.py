"""
Run once to generate sample inbound trade confirmation documents.
These simulate what a counterparty would send — actual values, minimal legal text.
    python sample_templates/generate_trade_confirmations.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = os.path.dirname(__file__)


def make_confirmation(filename, data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # Header bar — sender info
    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(2)
    r = header.add_run(data["sender_name"])
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    sub.add_run(f"{data['sender_address']}    |    {data['sender_email']}    |    {data['sender_phone']}")
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Divider
    doc.add_paragraph("─" * 80).paragraph_format.space_after = Pt(8)

    # Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(14)
    r = title.add_run("TRADE CONFIRMATION")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Reference line
    ref_line = doc.add_paragraph()
    ref_line.paragraph_format.space_after = Pt(14)
    ref_line.add_run(f"Confirmation ID: ").bold = True
    ref_line.add_run(f"{data['confirmation_id']}    ")
    ref_line.add_run("Trade Date: ").bold = True
    ref_line.add_run(f"{data['trade_date']}    ")
    ref_line.add_run("Settlement Date: ").bold = True
    ref_line.add_run(data['settlement_date'])
    for run in ref_line.runs:
        run.font.size = Pt(10.5)

    # Main details table
    def section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    def detail_table(rows):
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)
            row[0].paragraphs[0].runs[0].bold = True
            for cell in row:
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(2)
                    for run in para.runs:
                        run.font.size = Pt(10)
            # shade label column lightly
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tcPr = row[0]._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F5F5F5")
            tcPr.append(shd)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    section_heading("COUNTERPARTY DETAILS")
    detail_table([
        ("Counterparty", data["counterparty"]),
        ("Counterparty LEI", data["counterparty_lei"]),
        ("Broker / Desk", data["broker"]),
        ("Execution Venue", data["venue"]),
    ])

    section_heading("INSTRUMENT")
    detail_table([
        ("Security / Instrument", data["security"]),
        ("Ticker / Symbol", data["symbol"]),
        ("Asset Class", data["asset_class"]),
        ("ISIN", data.get("isin", "N/A")),
    ])

    section_heading("TRADE ECONOMICS")
    detail_table([
        ("Side",             data["side"]),
        ("Quantity",         data["quantity"]),
        ("Unit",             data["unit"]),
        ("Price",            data["price"]),
        ("Notional Value",   data["notional"]),
        ("Currency",         data["currency"]),
        ("Commission",       data.get("commission", "None")),
        ("Net Settlement",   data["net_settlement"]),
    ])

    if data.get("extra_rows"):
        section_heading("ADDITIONAL TERMS")
        detail_table(data["extra_rows"])

    # Footer note
    doc.add_paragraph("─" * 80).paragraph_format.space_before = Pt(14)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    note.add_run("Please review these details and confirm by return no later than ").font.size = Pt(9)
    r2 = note.add_run(data["confirm_by"])
    r2.bold = True; r2.font.size = Pt(9)
    r3 = note.add_run(". Any discrepancies must be reported to your designated relationship manager immediately.")
    r3.font.size = Pt(9)

    footer = doc.add_paragraph(
        f"This confirmation is issued by {data['sender_name']} and is subject to the terms of the "
        f"Master Trading Agreement in place between the parties. Ref: {data['mta_ref']}."
    )
    footer.paragraph_format.space_after = Pt(0)
    footer.runs[0].font.size = Pt(8.5)
    footer.runs[0].font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    path = os.path.join(OUT, filename)
    doc.save(path)
    print(f"Saved: {path}")


# ── Confirmation 1 — Equity block trade ──────────────────────
make_confirmation("trade_conf_equity_block.docx", {
    "sender_name":       "Harwick Capital Markets Ltd",
    "sender_address":    "25 Canary Wharf, London E14 5AB",
    "sender_email":      "confirmations@harwickcm.com",
    "sender_phone":      "+44 20 7946 0123",
    "confirmation_id":   "HCM-EQ-2026-004418",
    "trade_date":        "27 May 2026",
    "settlement_date":   "29 May 2026 (T+2)",
    "counterparty":      "Meridian Asset Management LP",
    "counterparty_lei":  "254900ABCD1234567890",
    "broker":            "Harwick Prime Brokerage",
    "venue":             "London Stock Exchange (XLON)",
    "security":          "Rolls-Royce Holdings plc — Ordinary Shares",
    "symbol":            "RR. LN",
    "asset_class":       "Equity",
    "isin":              "GB00B63H8491",
    "side":              "BUY",
    "quantity":          "500,000 shares",
    "unit":              "Shares",
    "price":             "GBP 4.82 per share",
    "notional":          "GBP 2,410,000.00",
    "currency":          "British Pound Sterling (GBP)",
    "commission":        "GBP 1,205.00 (0.05%)",
    "net_settlement":    "GBP 2,411,205.00",
    "confirm_by":        "28 May 2026 17:00 BST",
    "mta_ref":           "MTA-MAM-HCM-2024-001",
    "extra_rows": [
        ("Settlement Instructions", "CREST — Meridian Asset Mgmt Participant ID 98765"),
        ("Custodian",               "State Street Bank & Trust, London"),
        ("Fund / Account",          "Meridian European Opportunities Fund"),
    ],
})

# ── Confirmation 2 — FX Forward ──────────────────────────────
make_confirmation("trade_conf_fx_forward.docx", {
    "sender_name":       "Stellaris FX & Derivatives Ltd",
    "sender_address":    "10 Upper Bank Street, Canary Wharf, London E14 5NP",
    "sender_email":      "fx-confirms@stellarisfx.com",
    "sender_phone":      "+44 20 7946 0987",
    "confirmation_id":   "SFX-FWD-2026-001877",
    "trade_date":        "26 May 2026",
    "settlement_date":   "26 August 2026 (3M Forward)",
    "counterparty":      "Carrington Energy Partners LLC",
    "counterparty_lei":  "549300XYZW9988776655",
    "broker":            "Stellaris FX Desk",
    "venue":             "OTC Bilateral",
    "security":          "USD/EUR FX Forward",
    "symbol":            "USDEUR 3M FWD",
    "asset_class":       "FX Derivatives",
    "side":              "SELL USD / BUY EUR",
    "quantity":          "USD 10,000,000",
    "unit":              "USD Notional",
    "price":             "0.9245 (Forward Rate EUR/USD)",
    "notional":          "EUR 9,245,000.00",
    "currency":          "USD / EUR",
    "commission":        "None — spread included in rate",
    "net_settlement":    "EUR 9,245,000.00 delivered vs USD 10,000,000.00 received",
    "confirm_by":        "27 May 2026 12:00 BST",
    "mta_ref":           "ISDA-CEP-SFX-2025-003",
    "extra_rows": [
        ("Spot Rate at Trade",  "0.9198 EUR/USD"),
        ("Forward Points",      "+47 pips"),
        ("Fixing Source",       "ECB Reference Rate"),
        ("Settlement Method",   "Physical Delivery"),
        ("Netting",             "Subject to ISDA Master Agreement — gross settlement"),
    ],
})

# ── Confirmation 3 — Commodity Swap ──────────────────────────
make_confirmation("trade_conf_commodity_swap.docx", {
    "sender_name":       "Nexgen Commodity Brokers Inc.",
    "sender_address":    "1221 Avenue of the Americas, New York NY 10020",
    "sender_email":      "trade-confirms@nexgencb.com",
    "sender_phone":      "+1 212 555 0143",
    "confirmation_id":   "NCB-SWAP-2026-003341",
    "trade_date":        "22 May 2026",
    "settlement_date":   "Monthly cash settlement — first Business Day of following month",
    "counterparty":      "Starbridge Utilities Co. Ltd.",
    "counterparty_lei":  "213800LMNO4455667788",
    "broker":            "Nexgen Energy Derivatives Desk",
    "venue":             "OTC — CME Cleared Swap",
    "security":          "WTI Crude Oil Fixed-for-Floating Swap",
    "symbol":            "CL SWAP JUN26-DEC26",
    "asset_class":       "Commodity Derivatives",
    "isin":              "N/A — OTC Swap",
    "side":              "FIXED PAYER (Starbridge pays fixed, receives floating)",
    "quantity":          "10,000 barrels per month",
    "unit":              "Barrels (BBL)",
    "price":             "USD 78.50 per barrel (fixed leg)",
    "notional":          "USD 785,000 per month / USD 4,710,000 total",
    "currency":          "United States Dollar (USD)",
    "commission":        "USD 0.04 per barrel (broker commission)",
    "net_settlement":    "Cash settled monthly against NYMEX WTI front-month settlement",
    "confirm_by":        "23 May 2026 17:00 EST",
    "mta_ref":           "ISDA-SBU-NCB-2025-007",
    "extra_rows": [
        ("Floating Rate Index",   "NYMEX WTI Light Sweet Crude — front-month settlement price"),
        ("Calculation Period",    "June 2026 — December 2026 (7 months)"),
        ("Payment Frequency",     "Monthly — 1st Business Day following period end"),
        ("Clearing House",        "CME Clearing — account IDs on file"),
        ("Initial Margin",        "USD 210,000 posted at CME"),
        ("Variation Margin",      "Daily mark-to-market via CME"),
    ],
})

print("\nAll trade confirmations generated.")
