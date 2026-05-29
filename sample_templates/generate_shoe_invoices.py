"""
Generate sample shoe invoices for a running store.
Documents deliberately vary in layout, terminology, and style to stress-test LLM extraction.
    python sample_templates/generate_shoe_invoices.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.dirname(__file__)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


# ── Invoice 1 — Formal store invoice, Nike, table layout ──────────────────────
def make_invoice_1():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # Store header
    p = doc.add_paragraph()
    r = p.add_run("STRIDE & CO. RUNNING STORE")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph("42 Athlete's Way, Portland OR 97201  |  (503) 555-0198  |  orders@strideandco.com")
    sub.runs[0].font.size = Pt(9); sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub.paragraph_format.space_after = Pt(10)

    doc.add_paragraph("─" * 80).paragraph_format.space_after = Pt(8)

    title = doc.add_paragraph()
    r = title.add_run("PURCHASE INVOICE")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    title.paragraph_format.space_after = Pt(14)

    # Invoice meta
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    meta.add_run("Invoice No: ").bold = True
    meta.add_run("INV-2026-00841    ")
    meta.add_run("Date: ").bold = True
    meta.add_run("22 May 2026    ")
    meta.add_run("Confirmation Ref: ").bold = True
    meta.add_run("CONF-Nike-PegPrem-0841")
    for run in meta.runs:
        run.font.size = Pt(10.5)

    # Customer
    cust = doc.add_table(rows=0, cols=2)
    cust.style = "Table Grid"
    for lbl, val in [
        ("Customer Name", "Jordan Whitfield"),
        ("Phone", "(503) 555-7723"),
        ("Email", "jordan.whitfield@email.com"),
        ("Delivery Address", "14 Maple Drive, Portland OR 97210"),
    ]:
        row = cust.add_row().cells
        row[0].text = lbl; row[1].text = val
        row[0].paragraphs[0].runs[0].bold = True
        shade_cell(row[0], "F5F5F5")
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Product details
    p2 = doc.add_paragraph()
    r2 = p2.add_run("PRODUCT DETAILS")
    r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p2.paragraph_format.space_before = Pt(8); p2.paragraph_format.space_after = Pt(4)

    prod = doc.add_table(rows=0, cols=2)
    prod.style = "Table Grid"
    for lbl, val in [
        ("Brand",              "Nike"),
        ("Model",              "Pegasus Premium"),
        ("Colorway",           "Thunder Blue / White / Volt"),
        ("Size (US Men's)",    "10.5"),
        ("Stack Height",       "37 mm"),
        ("Category",           "Daily Trainer / Road Running"),
        ("Unit Price",         "USD 160.00"),
        ("Qty",                "1"),
        ("Total",              "USD 160.00"),
    ]:
        row = prod.add_row().cells
        row[0].text = lbl; row[1].text = val
        row[0].paragraphs[0].runs[0].bold = True
        shade_cell(row[0], "F5F5F5")
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Footer
    doc.add_paragraph("─" * 80).paragraph_format.space_before = Pt(10)
    f = doc.add_paragraph("Thank you for shopping at Stride & Co. All sales are final. Returns accepted within 30 days with receipt.")
    f.runs[0].font.size = Pt(9); f.runs[0].font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    doc.save(os.path.join(OUT, "shoe_invoice_nike_pegasus.docx"))
    print("Saved: shoe_invoice_nike_pegasus.docx")


# ── Invoice 2 — Casual receipt-style, Asics, minimal table ───────────────────
def make_invoice_2():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Pacers Running — Sales Receipt")
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x00, 0x4e, 0x9a)
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph("88 Finish Line Blvd, Seattle WA 98101  •  pacersrunning.com")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9); sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub.paragraph_format.space_after = Pt(16)

    def kv(doc, label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True; r1.font.size = Pt(10.5)
        r2 = p.add_run(value)
        r2.font.size = Pt(10.5)

    kv(doc, "Receipt No", "PAC-REC-20260519-3347")
    kv(doc, "Transaction Date", "19 May 2026")
    kv(doc, "Sales Associate", "Taylor M.")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    kv(doc, "Customer", "Sam Okafor")
    kv(doc, "Confirmation Number", "CONF-Asics-KayanoV2-3347")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Item Purchased")
    r3.bold = True; r3.font.size = Pt(11)
    p3.paragraph_format.space_after = Pt(6)

    # Simple list style
    items = [
        ("Shoe Brand",       "Asics"),
        ("Model Name",       "Gel-Kayano 31"),
        ("Colour",           "Midnight / Pure Silver"),
        ("US Size",          "9"),
        ("Stack Height (mm)","40"),
        ("Drop",             "10 mm"),
        ("Price Paid",       "$165.00 USD"),
        ("Payment Method",   "Visa ending 4872"),
    ]
    for lbl, val in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        r1 = p.add_run(f"{lbl}: ")
        r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    note = doc.add_paragraph("Please retain this receipt for warranty and return purposes. Exchanges within 45 days.")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(9); note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(os.path.join(OUT, "shoe_invoice_asics_kayano.docx"))
    print("Saved: shoe_invoice_asics_kayano.docx")


# ── Invoice 3 — B2B wholesale order, Adidas, multiple lines ──────────────────
def make_invoice_3():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    hdr = doc.add_paragraph()
    r = hdr.add_run("Fleet Foot Wholesale Distributors")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    hdr.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("1100 Commerce Park, Chicago IL 60607  |  wholesale@fleetfoot.com  |  +1 312 555 0220").runs[0].font.size = Pt(9)
    doc.add_paragraph("─" * 80).paragraph_format.space_after = Pt(6)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_p.add_run("WHOLESALE PURCHASE ORDER — FOOTWEAR")
    t.bold = True; t.font.size = Pt(12); t.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    title_p.paragraph_format.space_after = Pt(14)

    # PO meta table
    meta_tbl = doc.add_table(rows=0, cols=4)
    meta_tbl.style = "Table Grid"
    labels = ["PO Number", "FF-PO-2026-0048", "Order Date", "14 May 2026"]
    row = meta_tbl.add_row().cells
    for i, txt in enumerate(labels):
        row[i].text = txt
        if i % 2 == 0:
            row[i].paragraphs[0].runs[0].bold = True
            shade_cell(row[i], "ECECEC")
        for para in row[i].paragraphs:
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(10)

    labels2 = ["Buyer / Store", "Metro Run Club, Nashville TN", "Confirmation Ref", "CONF-Adidas-Boston14-0048"]
    row2 = meta_tbl.add_row().cells
    for i, txt in enumerate(labels2):
        row2[i].text = txt
        if i % 2 == 0:
            row2[i].paragraphs[0].runs[0].bold = True
            shade_cell(row2[i], "ECECEC")
        for para in row2[i].paragraphs:
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    p_head = doc.add_paragraph()
    p_head.add_run("Line Item Detail").bold = True
    p_head.runs[0].font.size = Pt(11)
    p_head.paragraph_format.space_after = Pt(4)

    # Line items table
    line_tbl = doc.add_table(rows=1, cols=6)
    line_tbl.style = "Table Grid"
    headers = ["Brand", "Model", "Colour / Colorway", "Size (EU)", "Stack Ht.", "Unit Cost"]
    hrow = line_tbl.rows[0].cells
    for i, h in enumerate(headers):
        hrow[i].text = h
        hrow[i].paragraphs[0].runs[0].bold = True
        shade_cell(hrow[i], "1A1A2E")
        hrow[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        hrow[i].paragraphs[0].runs[0].font.size = Pt(9)

    line_data = [
        ("Adidas", "Boston 14", "Core Black / Cloud White", "43 (US 9.5)", "39 mm", "$140.00"),
        ("Adidas", "Boston 14", "Wonder White / Spark",     "44 (US 10)",  "39 mm", "$140.00"),
        ("Adidas", "Boston 14", "Lucid Blue / Spark",       "42 (US 9)",   "39 mm", "$140.00"),
    ]
    for brand, model, color, size, stack, price in line_data:
        r = line_tbl.add_row().cells
        for i, val in enumerate([brand, model, color, size, stack, price]):
            r[i].text = val
            for para in r[i].paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    totals = doc.add_paragraph()
    totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals.add_run("Subtotal: $420.00   |   Tax (8.5%): $35.70   |   ").font.size = Pt(10)
    r_total = totals.add_run("Total Due: $455.70")
    r_total.bold = True; r_total.font.size = Pt(11)

    doc.add_paragraph("─" * 80).paragraph_format.space_before = Pt(10)
    footer = doc.add_paragraph("Payment due Net 30 from invoice date. Wire transfer or ACH accepted. Ref PO number on remittance.")
    footer.runs[0].font.size = Pt(9); footer.runs[0].font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    doc.save(os.path.join(OUT, "shoe_invoice_adidas_boston.docx"))
    print("Saved: shoe_invoice_adidas_boston.docx")


# ── Invoice 4 — Sparse email-style, Brooks, minimal formatting ───────────────
def make_invoice_4():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_paragraph("From: shop@runcityruns.com")
    doc.add_paragraph("To: alex.thornton@gmail.com")
    doc.add_paragraph("Subject: Your Order Confirmation — Brooks Ghost Max 2")
    doc.add_paragraph("Date: 27 May 2026")
    doc.add_paragraph()

    doc.add_paragraph("Hi Alex,")
    doc.add_paragraph()
    body = doc.add_paragraph(
        "Thanks for your order at Run City! Here's a summary of what we've got lined up for you:"
    )
    body.paragraph_format.space_after = Pt(10)

    # Sparse details — label: value format, no table
    details = [
        ("Order Confirmation #", "CONF-Brooks-GhostMax2-7781"),
        ("Order Date",           "27 May 2026"),
        ("Item",                 "Brooks Ghost Max 2"),
        ("Brand",                "Brooks"),
        ("Shoe Colorway",        "Oyster / Chateau Grey / Ebony"),
        ("Men's US Size",        "11"),
        ("Stack Height",         "46mm"),
        ("Price",                "$165.00"),
        ("Shipping",             "$0.00 — free standard"),
        ("Total Charged",        "$165.00"),
        ("Est. Delivery",        "2–4 business days"),
    ]
    for lbl, val in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{lbl}: ")
        r1.bold = True
        p.add_run(val)

    doc.add_paragraph()
    sign = doc.add_paragraph(
        "If you have any questions, reply to this email or call us at (206) 555-0134. "
        "Your shoes will ship within 1 business day."
    )
    sign.paragraph_format.space_after = Pt(10)
    doc.add_paragraph("Best,")
    doc.add_paragraph("The Run City Team")

    doc.save(os.path.join(OUT, "shoe_invoice_brooks_ghostmax.docx"))
    print("Saved: shoe_invoice_brooks_ghostmax.docx")


make_invoice_1()
make_invoice_2()
make_invoice_3()
make_invoice_4()
print("\nAll shoe invoices generated.")
