"""Run this script once to generate sample_contract.docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5) if level == 1 else Pt(10.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def clause(number, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"{number}.  {title}.  ")
    run.bold = True
    run.font.size = Pt(10.5)
    p.add_run(text)

# ── Title block ───────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("MASTER SUPPLY AND SERVICES AGREEMENT")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Contract Reference: ").bold = True
meta.add_run("MSA-2026-0047    ")
meta.add_run("Effective Date: ").bold = True
meta.add_run("1 June 2026    ")
meta.add_run("Governing Law: ").bold = True
meta.add_run("England & Wales")

doc.add_paragraph()

# ── Parties ───────────────────────────────────────────────────
heading("PARTIES")
body(
    "This Master Supply and Services Agreement (\"Agreement\") is entered into as of the Effective Date "
    "set forth above by and between:"
)
body(
    "Nortech Commodities Limited, a company incorporated under the laws of England and Wales with "
    "registered number 09471823 and having its registered office at 14 Moorgate, London EC2R 6DA "
    "(\"Supplier\"); and"
)
body(
    "Carrington Energy Partners LLC, a limited liability company organised under the laws of the State "
    "of Delaware, United States, with its principal place of business at 350 Fifth Avenue, New York, "
    "NY 10118 (\"Buyer\")."
)
body("Supplier and Buyer are each referred to herein individually as a \"Party\" and collectively as the \"Parties\".")

# ── Recitals ─────────────────────────────────────────────────
heading("RECITALS")
body(
    "WHEREAS, Supplier is engaged in the business of supplying, trading and delivering physical commodities "
    "and related ancillary services; and"
)
body(
    "WHEREAS, Buyer desires to purchase from Supplier, and Supplier desires to sell and deliver to Buyer, "
    "such commodities and services on the terms and conditions set forth herein."
)
body("NOW, THEREFORE, in consideration of the mutual covenants and agreements hereinafter set forth and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:")

# ── Section 1 ────────────────────────────────────────────────
heading("SECTION 1 — DEFINITIONS AND INTERPRETATION")
clause("1.1", "Definitions",
    "\"Affiliate\" means, with respect to any Person, any other Person that directly or indirectly, "
    "through one or more intermediaries, Controls, is Controlled by, or is under common Control with, "
    "such Person. \"Business Day\" means any day other than a Saturday, Sunday or public holiday in "
    "England and Wales. \"Confidential Information\" means all non-public information disclosed by one "
    "Party to the other, whether orally, in writing, or by any other means, that is designated as "
    "confidential or that reasonably should be understood to be confidential given the nature of the "
    "information and the circumstances of disclosure.")
clause("1.2", "Interpretation",
    "In this Agreement, unless the context otherwise requires: (i) references to a statute or statutory "
    "provision include any subordinate legislation made under it and any amendment, re-enactment or "
    "replacement of it; (ii) the singular includes the plural and vice versa; (iii) a reference to "
    "\"including\" shall be construed as \"including without limitation\"; and (iv) headings are inserted "
    "for convenience only and shall not affect the construction of this Agreement.")

# ── Section 2 ────────────────────────────────────────────────
heading("SECTION 2 — TERM AND TERMINATION")
clause("2.1", "Term",
    "This Agreement shall commence on the Effective Date and shall continue in full force and effect "
    "for an initial period of thirty-six (36) months (the \"Initial Term\"), unless earlier terminated "
    "in accordance with the provisions hereof. Thereafter, this Agreement shall automatically renew for "
    "successive periods of twelve (12) months each (each a \"Renewal Term\") unless either Party provides "
    "the other with not less than ninety (90) days' prior written notice of its intention not to renew.")
clause("2.2", "Termination for Cause",
    "Either Party may terminate this Agreement immediately upon written notice if the other Party: "
    "(a) commits a material breach of any provision of this Agreement and fails to cure such breach "
    "within thirty (30) days after receipt of written notice describing the breach in reasonable detail; "
    "(b) becomes Insolvent; or (c) ceases or threatens to cease to carry on business.")
clause("2.3", "Termination for Convenience",
    "After the expiry of the Initial Term, either Party may terminate this Agreement for convenience by "
    "providing not less than one hundred and eighty (180) days' prior written notice to the other Party.")

# ── Page break ───────────────────────────────────────────────
doc.add_page_break()

# ── Section 3 — Commercial Details ───────────────────────────
heading("SECTION 3 — COMMERCIAL DETAILS")
body(
    "The following commercial terms govern each Transaction entered into under this Agreement. "
    "Individual Transactions may be confirmed by way of a Transaction Confirmation substantially in "
    "the form set out in Schedule 1. In the event of any conflict between a Transaction Confirmation "
    "and this Agreement, the Transaction Confirmation shall prevail in respect of that Transaction only."
)

# commercial table
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Commercial Term"
hdr[1].text = "Detail"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(2)
    cell._tc.get_or_add_tcPr()

rows = [
    ("Counterparty (Buyer)",       "Carrington Energy Partners LLC"),
    ("Counterparty (Supplier)",    "Nortech Commodities Limited"),
    ("Commodity",                  "Natural Gas — TTF Day-Ahead"),
    ("Contract Volume",            "50,000 MMBtu per month"),
    ("Minimum Off-Take Volume",    "40,000 MMBtu per month"),
    ("Maximum Off-Take Volume",    "60,000 MMBtu per month"),
    ("Unit Price",                 "USD 8.45 per MMBtu"),
    ("Price Basis",                "Fixed — not subject to index adjustment"),
    ("Currency",                   "United States Dollar (USD)"),
    ("Delivery Point",             "Title Transfer Facility (TTF), Netherlands"),
    ("Delivery Period Start",      "1 July 2026"),
    ("Delivery Period End",        "30 June 2027"),
    ("Payment Terms",              "Net 30 days from invoice date"),
    ("Invoice Currency",           "USD"),
    ("Late Payment Interest",      "SONIA + 2.00% per annum"),
    ("Credit Support",             "Parent Guarantee — Carrington Global Holdings Inc."),
    ("Credit Limit",               "USD 5,000,000"),
]
for label, value in rows:
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value
    for cell in row:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

clause("3.1", "Price Adjustment",
    "The Unit Price specified above is fixed for the Delivery Period and shall not be subject to "
    "adjustment by reference to any index, market price or other external benchmark, except as "
    "expressly agreed in writing by the Parties by way of an executed Amendment to this Agreement.")
clause("3.2", "Invoicing",
    "Supplier shall issue invoices to Buyer on a monthly basis, no later than the fifth (5th) Business "
    "Day following the end of each calendar month in which delivery has occurred. Each invoice shall "
    "set out in reasonable detail: (i) the volume of Commodity delivered; (ii) the applicable Unit "
    "Price; (iii) the total amount due; and (iv) applicable taxes and levies, if any.")
clause("3.3", "Taxes",
    "Each Party shall be responsible for all taxes, duties, levies and other governmental charges "
    "imposed on it in connection with this Agreement. All amounts stated in this Agreement are "
    "exclusive of value added tax (\"VAT\") or any equivalent tax, which shall be added at the "
    "applicable rate where required by law.")

# ── Section 4 ────────────────────────────────────────────────
heading("SECTION 4 — REPRESENTATIONS AND WARRANTIES")
clause("4.1", "Mutual Representations",
    "Each Party represents and warrants to the other Party as of the Effective Date and as of the "
    "date of each Transaction that: (a) it is duly organised, validly existing and in good standing "
    "under the laws of its jurisdiction of organisation; (b) it has full power and authority to "
    "execute, deliver and perform its obligations under this Agreement; (c) this Agreement has been "
    "duly authorised, executed and delivered by it and constitutes its legal, valid and binding "
    "obligation enforceable against it in accordance with its terms; and (d) the execution, delivery "
    "and performance of this Agreement do not violate any applicable law, regulation or order.")

# ── Section 5 ────────────────────────────────────────────────
heading("SECTION 5 — LIMITATION OF LIABILITY AND INDEMNIFICATION")
clause("5.1", "Limitation of Liability",
    "To the maximum extent permitted by applicable law, neither Party shall be liable to the other "
    "for any indirect, incidental, special, consequential, exemplary or punitive damages, including "
    "loss of profits, loss of revenue, loss of business or loss of data, arising out of or in "
    "connection with this Agreement, even if such Party has been advised of the possibility of such "
    "damages. Each Party's aggregate liability under this Agreement shall not exceed the total fees "
    "paid or payable in the twelve (12) month period immediately preceding the event giving rise to "
    "the claim.")
clause("5.2", "Indemnification",
    "Each Party (\"Indemnifying Party\") shall defend, indemnify and hold harmless the other Party "
    "and its Affiliates, officers, directors, employees and agents (collectively, \"Indemnified Parties\") "
    "from and against any and all claims, damages, losses, costs and expenses (including reasonable "
    "legal fees) arising out of or resulting from: (a) any breach by the Indemnifying Party of its "
    "representations, warranties or obligations under this Agreement; or (b) the gross negligence or "
    "wilful misconduct of the Indemnifying Party.")

# ── Section 6 ────────────────────────────────────────────────
heading("SECTION 6 — GENERAL PROVISIONS")
clause("6.1", "Governing Law",
    "This Agreement and any dispute or claim arising out of or in connection with it or its subject "
    "matter or formation (including non-contractual disputes or claims) shall be governed by and "
    "construed in accordance with the laws of England and Wales.")
clause("6.2", "Dispute Resolution",
    "Any dispute arising out of or in connection with this Agreement, including any question regarding "
    "its existence, validity or termination, shall be referred to and finally resolved by arbitration "
    "under the LCIA Rules, which Rules are deemed to be incorporated by reference into this clause. "
    "The seat of arbitration shall be London. The language of the arbitration shall be English. "
    "The number of arbitrators shall be three (3).")
clause("6.3", "Entire Agreement",
    "This Agreement, together with all Schedules and Transaction Confirmations, constitutes the entire "
    "agreement between the Parties with respect to the subject matter hereof and supersedes all prior "
    "and contemporaneous agreements, understandings, negotiations and discussions, whether oral or "
    "written, of the Parties.")
clause("6.4", "Amendments",
    "No amendment to this Agreement shall be effective unless it is in writing and signed by duly "
    "authorised representatives of both Parties.")
clause("6.5", "Severability",
    "If any provision of this Agreement is held to be invalid, illegal or unenforceable in any "
    "respect, such invalidity, illegality or unenforceability shall not affect any other provision "
    "hereof, and this Agreement shall be construed as if such invalid, illegal or unenforceable "
    "provision had never been contained herein.")

doc.add_paragraph()

# ── Signature block ───────────────────────────────────────────
heading("EXECUTION")
body("IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.")
doc.add_paragraph()

sig_table = doc.add_table(rows=6, cols=2)
sig_data = [
    ("NORTECH COMMODITIES LIMITED", "CARRINGTON ENERGY PARTNERS LLC"),
    ("", ""),
    ("Signature: _______________________", "Signature: _______________________"),
    ("Name:      _______________________", "Name:      _______________________"),
    ("Title:     _______________________", "Title:     _______________________"),
    ("Date:      _______________________", "Date:      _______________________"),
]
for i, (left, right) in enumerate(sig_data):
    row = sig_table.rows[i].cells
    row[0].text = left
    row[1].text = right
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = (i == 0)
                run.font.size = Pt(10)

out_path = os.path.join(os.path.dirname(__file__), "sample_contract.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
