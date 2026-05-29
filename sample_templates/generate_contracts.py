"""
Run once to generate three sample contracts with different CPs, commodities and commercial terms.
    python sample_templates/generate_contracts.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = os.path.dirname(__file__)


def make_contract(filename, ref, effective_date, supplier, buyer, commodity,
                  volume, min_vol, max_vol, unit_price, price_basis, currency,
                  delivery_point, period_start, period_end, payment_terms,
                  late_interest, credit_support, credit_limit, governing_law,
                  term_months):

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def heading(text, level=1):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11.5) if level == 1 else Pt(10.5)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def clause(num, title, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"{num}.  {title}.  ")
        r.bold = True
        r.font.size = Pt(10.5)
        p.add_run(text)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MASTER SUPPLY AND SERVICES AGREEMENT")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Contract Reference: ").bold = True
    meta.add_run(f"{ref}    ")
    meta.add_run("Effective Date: ").bold = True
    meta.add_run(f"{effective_date}    ")
    meta.add_run("Governing Law: ").bold = True
    meta.add_run(governing_law)
    doc.add_paragraph()

    heading("PARTIES")
    body(f'This Master Supply and Services Agreement ("Agreement") is entered into as of the Effective Date by and between:')
    body(f'{supplier} ("Supplier"); and')
    body(f'{buyer} ("Buyer").')
    body('Supplier and Buyer are each referred to herein individually as a "Party" and collectively as the "Parties".')

    heading("RECITALS")
    body("WHEREAS, Supplier is engaged in the business of supplying, trading and delivering physical commodities and related services; and")
    body("WHEREAS, Buyer desires to purchase from Supplier, and Supplier desires to sell to Buyer, such commodities on the terms herein.")
    body("NOW, THEREFORE, in consideration of the mutual covenants herein and other good and valuable consideration, the Parties agree as follows:")

    heading("SECTION 1 — DEFINITIONS AND INTERPRETATION")
    clause("1.1", "Definitions",
        '"Affiliate" means any Person that directly or indirectly Controls, is Controlled by, or is under '
        'common Control with such Person. "Business Day" means any day other than a Saturday, Sunday or '
        'public holiday in the applicable jurisdiction. "Confidential Information" means all non-public '
        'information disclosed by one Party to the other that is designated as confidential or that '
        'reasonably should be understood to be confidential.')
    clause("1.2", "Interpretation",
        "Unless context otherwise requires: (i) references to a statute include any subordinate legislation "
        "and any amendment or replacement thereof; (ii) the singular includes the plural and vice versa; "
        "(iii) 'including' shall mean 'including without limitation'; and (iv) headings are for convenience only.")

    heading("SECTION 2 — TERM AND TERMINATION")
    clause("2.1", "Initial Term",
        f"This Agreement commences on the Effective Date and continues for {term_months} months (the "
        '"Initial Term"), unless earlier terminated. Thereafter it renews automatically for successive '
        "12-month periods unless either Party gives not less than 90 days' prior written notice.")
    clause("2.2", "Termination for Cause",
        "Either Party may terminate immediately on written notice if the other: (a) materially breaches "
        "this Agreement and fails to cure within 30 days of written notice; (b) becomes insolvent or "
        "subject to insolvency proceedings; or (c) ceases to carry on business.")
    clause("2.3", "Termination for Convenience",
        "After the Initial Term, either Party may terminate for convenience on not less than 180 days' prior written notice.")

    doc.add_page_break()

    heading("SECTION 3 — COMMERCIAL DETAILS")
    body(
        "The following commercial terms govern each Transaction under this Agreement. "
        "Individual Transactions may be confirmed by a Transaction Confirmation in the form set out in Schedule 1. "
        "In the event of conflict, the Transaction Confirmation shall prevail for that Transaction."
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Commercial Term"
    hdr[1].text = "Detail"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(2)

    rows = [
        ("Counterparty (Buyer)", buyer),
        ("Counterparty (Supplier)", supplier),
        ("Commodity", commodity),
        ("Contract Volume", volume),
        ("Minimum Off-Take Volume", min_vol),
        ("Maximum Off-Take Volume", max_vol),
        ("Unit Price", unit_price),
        ("Price Basis", price_basis),
        ("Currency", currency),
        ("Delivery Point", delivery_point),
        ("Delivery Period Start", period_start),
        ("Delivery Period End", period_end),
        ("Payment Terms", payment_terms),
        ("Late Payment Interest", late_interest),
        ("Credit Support", credit_support),
        ("Credit Limit", credit_limit),
    ]
    for lbl, val in rows:
        row = table.add_row().cells
        row[0].text = lbl; row[1].text = val
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    clause("3.1", "Price Adjustment",
        f"The Unit Price of {unit_price} is fixed for the Delivery Period and shall not be subject to "
        "adjustment by reference to any index or external benchmark except as expressly agreed in writing.")
    clause("3.2", "Invoicing",
        "Supplier shall issue invoices monthly, no later than the 5th Business Day following month-end. "
        "Each invoice shall state: (i) volume delivered; (ii) applicable Unit Price; (iii) total amount due; "
        "and (iv) applicable taxes.")
    clause("3.3", "Taxes",
        "Each Party bears its own taxes. All amounts are exclusive of VAT or equivalent, which shall be "
        "added at the applicable rate where required by law.")

    heading("SECTION 4 — REPRESENTATIONS AND WARRANTIES")
    clause("4.1", "Mutual Representations",
        "Each Party represents and warrants that: (a) it is duly organised and validly existing; "
        "(b) it has full authority to execute and perform this Agreement; (c) this Agreement is its "
        "legal, valid and binding obligation; and (d) execution does not violate any applicable law.")

    heading("SECTION 5 — LIMITATION OF LIABILITY")
    clause("5.1", "Exclusion of Consequential Loss",
        "Neither Party shall be liable for any indirect, incidental, consequential or punitive damages. "
        "Aggregate liability shall not exceed total fees paid in the 12 months preceding the claim.")
    clause("5.2", "Indemnification",
        "Each Party shall indemnify the other against claims arising from its own breach, gross negligence "
        "or wilful misconduct.")

    heading("SECTION 6 — GENERAL PROVISIONS")
    clause("6.1", "Governing Law", f"This Agreement is governed by the laws of {governing_law}.")
    clause("6.2", "Dispute Resolution",
        "Disputes shall be resolved by arbitration under LCIA Rules, seated in London, in English, "
        "before three arbitrators.")
    clause("6.3", "Entire Agreement",
        "This Agreement constitutes the entire agreement and supersedes all prior understandings.")
    clause("6.4", "Amendments", "Amendments are only effective if in writing and signed by both Parties.")
    clause("6.5", "Severability",
        "If any provision is held invalid, the remaining provisions continue in full force.")

    doc.add_paragraph()
    heading("EXECUTION")
    body("IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()

    sig = doc.add_table(rows=6, cols=2)
    sig_data = [
        (supplier.split(",")[0].upper(), buyer.split(",")[0].upper()),
        ("", ""),
        ("Signature: _______________________", "Signature: _______________________"),
        ("Name:      _______________________", "Name:      _______________________"),
        ("Title:     _______________________", "Title:     _______________________"),
        ("Date:      _______________________", "Date:      _______________________"),
    ]
    for i, (left, right) in enumerate(sig_data):
        row = sig.rows[i].cells
        row[0].text = left; row[1].text = right
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = (i == 0); run.font.size = Pt(10)

    path = os.path.join(OUT, filename)
    doc.save(path)
    print(f"Saved: {path}")


# ── Contract 1 — Natural Gas, Fixed Price, England & Wales ───
make_contract(
    filename="contract_nortech_carrington.docx",
    ref="MSA-2026-0047",
    effective_date="1 June 2026",
    supplier="Nortech Commodities Limited, incorporated in England & Wales (No. 09471823), 14 Moorgate, London EC2R 6DA",
    buyer="Carrington Energy Partners LLC, a Delaware LLC, 350 Fifth Avenue, New York NY 10118",
    commodity="Natural Gas — TTF Day-Ahead",
    volume="50,000 MMBtu per month",
    min_vol="40,000 MMBtu per month",
    max_vol="60,000 MMBtu per month",
    unit_price="USD 8.45 per MMBtu",
    price_basis="Fixed — not subject to index adjustment",
    currency="United States Dollar (USD)",
    delivery_point="Title Transfer Facility (TTF), Netherlands",
    period_start="1 July 2026",
    period_end="30 June 2027",
    payment_terms="Net 30 days from invoice date",
    late_interest="SONIA + 2.00% per annum",
    credit_support="Parent Guarantee — Carrington Global Holdings Inc.",
    credit_limit="USD 5,000,000",
    governing_law="England and Wales",
    term_months=36,
)

# ── Contract 2 — Power, Floating Price, New York ─────────────
make_contract(
    filename="contract_voltex_meridian.docx",
    ref="PWR-2026-0112",
    effective_date="15 March 2026",
    supplier="Voltex Power Trading GmbH, registered in Germany (HRB 204831), Taunusanlage 8, 60329 Frankfurt",
    buyer="Meridian Industrial Corp., a New York corporation, 1 Penn Plaza, New York NY 10119",
    commodity="Electricity — PJM Western Hub Peak",
    volume="25 MW per hour during Peak Hours",
    min_vol="20 MW per hour",
    max_vol="30 MW per hour",
    unit_price="USD 52.00 per MWh (subject to Section 3.1 index adjustment)",
    price_basis="Floating — indexed to PJM Day-Ahead LMP monthly average",
    currency="United States Dollar (USD)",
    delivery_point="PJM Western Hub",
    period_start="1 April 2026",
    period_end="31 March 2027",
    payment_terms="Net 20 days from invoice date",
    late_interest="Fed Funds Rate + 3.00% per annum",
    credit_support="Letter of Credit — Deutsche Bank AG",
    credit_limit="USD 2,500,000",
    governing_law="State of New York, United States",
    term_months=24,
)

# ── Contract 3 — LNG, Spot, Singapore ───────────────────────
make_contract(
    filename="contract_pacificlng_starbridge.docx",
    ref="LNG-2026-0389",
    effective_date="1 January 2026",
    supplier="Pacific LNG Trading Pte. Ltd., incorporated in Singapore (UEN 201934782K), 8 Marina Boulevard, Singapore 018981",
    buyer="Starbridge Utilities Co. Ltd., a Hong Kong company, 88 Queensway, Admiralty, Hong Kong",
    commodity="Liquefied Natural Gas (LNG) — DES Delivered",
    volume="1 LNG Cargo per month (approx. 65,000 m³)",
    min_vol="0.8 LNG Cargo per month",
    max_vol="1.2 LNG Cargoes per month",
    unit_price="USD 11.20 per MMBtu",
    price_basis="Fixed slope linked to JKM — 12.5% x JKM Platts assessment",
    currency="United States Dollar (USD)",
    delivery_point="Starbridge FSRU Terminal, Lantau Island, Hong Kong",
    period_start="1 February 2026",
    period_end="31 January 2028",
    payment_terms="5 Business Days after Bill of Lading date",
    late_interest="USD SOFR + 2.50% per annum",
    credit_support="Standby Letter of Credit — HSBC Hong Kong",
    credit_limit="USD 8,000,000",
    governing_law="Singapore",
    term_months=24,
)

print("\nAll contracts generated.")
