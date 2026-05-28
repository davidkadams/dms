from docx import Document
import io


def extract_text(content: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "docx":
        return _extract_docx(content)
    if ext in ("txt", "md"):
        return content.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: .{ext}")


def _extract_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)
